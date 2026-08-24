from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_experiment_report():
    # 创建文档对象
    doc = Document()
    
    # 设置中文字体支持
    doc.styles['Normal'].font.name = u'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # --- 标题 ---
    title = doc.add_heading('实验报告', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- 个人信息表格 ---
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    
    # 填充表格内容
    data = [
        ('学 号', '24020007135', '姓 名', '吴虹霖'),
        ('专业班级', '________', '课程名称', '离散数学II'),
        ('学 期', '2024-2025学年', '任课教师', '________'),
        ('完成日期', '2024年__月__日', '上机时间', '周__ 第__节'),
    ]
    
    for i, row_data in enumerate(data):
        cells = table.rows[i].cells
        cells[0].text = row_data[0]
        cells[1].text = row_data[1]
        cells[2].text = row_data[2]
        cells[3].text = row_data[3]

    # 实验名称行
    row = table.add_row()
    row.cells[0].text = '实验名称'
    # 合并剩余单元格
    row.cells[1].merge(row.cells[2]).merge(row.cells[3])
    row.cells[1].text = '图的可视化表示，矩阵表示，基本要素分析，连通性判定'

    doc.add_paragraph() # 空行

    # --- 一、实验要求 ---
    h1 = doc.add_heading('一、实验要求', level=2)
    requirements = [
        "图的可视化表示：根据无向图的关联矩阵，绘制图的结构，清晰标注顶点 vi 和边 ej。",
        "矩阵转换：根据关联矩阵，计算并输出该图的邻接矩阵 A。",
        "基本要素分析：统计并输出每个顶点的度数（Degree）及邻域（Neighborhood）。",
        "规范性判断：判断该无向图是否为“简单图”（即不存在自环和多重边）。",
        "连通性判定：分析图的连通性，输出连通分支的数量及每个分支包含的顶点集合。"
    ]
    for req in requirements:
        p = doc.add_paragraph(req, style='List Number')

    # --- 二、实验内容及步骤 ---
    doc.add_heading('二、实验内容及步骤', level=2)
    
    doc.add_heading('1. 需求与概要设计', level=3)
    p = doc.add_paragraph()
    p.add_run('输入：').bold = True
    p.add_run('一个 n x m 的关联矩阵（Incidence Matrix）。\n')
    p.add_run('处理逻辑：').bold = True
    p.add_run('利用 NumPy 遍历关联矩阵构建邻接矩阵；利用 NetworkX 构建图对象并计算连通分量；通过 Matplotlib 进行可视化（针对多重边进行弧度优化）。')

    doc.add_heading('2. 详细代码实现', level=3)
    doc.add_paragraph('本实验核心代码如下（基于Python 3.x）：')
    
    code_content = """
import numpy as np
import networkx as nx
import matplotlib.pyplot as plt

# 核心模块：关联矩阵转邻接矩阵
def get_adjacency_from_incidence(inc_matrix):
    rows, cols = inc_matrix.shape
    adj_matrix = np.zeros((rows, rows), dtype=int)
    for j in range(cols):
        nodes = np.where(inc_matrix[:, j] != 0)[0]
        if len(nodes) == 2:
            u, v = nodes
            adj_matrix[u][v] += 1
            adj_matrix[v][u] += 1
        elif len(nodes) == 1:
            u = nodes[0]
            adj_matrix[u][u] += 1 
    return adj_matrix

# 核心模块：简单图判定
def check_simple_graph(adj_matrix):
    has_loops = np.any(np.diagonal(adj_matrix) > 0)
    has_parallel_edges = np.any(adj_matrix > 1)
    return not (has_loops or has_parallel_edges)
"""
    doc.add_paragraph(code_content, style='Quote')

    doc.add_heading('3. 实验结果分析', level=3)
    doc.add_paragraph('选取代码中定义的矩阵 M3（包含多重边）作为测试用例：')
    
    p = doc.add_paragraph()
    p.add_run('程序运行输出结果：').bold = True
    result_text = """
--- 邻接矩阵 ---
[[0 2 0 1]
 [2 0 1 1]
 [0 1 0 0]
 [1 1 0 0]]
(分析：v1和v2之间值为2，表示存在多重边)

--- 简单图判断 ---
是否存在自环: False, 是否存在多重边: True
结论: 该图 不是 简单图

--- 连通性分析 ---
连通分支数量: 1 (所有节点连通)
"""
    doc.add_paragraph(result_text, style='No Spacing')
    
    doc.add_paragraph('可视化结果说明：程序生成的图形清晰展示了v1与v2之间的两条弯曲连线，验证了对多重边的处理逻辑。')

    # --- 三、心得总结 ---
    doc.add_heading('三、心得总结', level=2)
    
    summary = [
        "遇到的问题：在使用 Matplotlib 绘制多重图时，发现默认设置会将多条边重叠绘制成一条直线，导致视觉信息丢失。",
        "解决方法：引入了 connectionstyle='arc3, rad=...' 参数。通过判断两点间的边数，为每一条边动态分配不同的弧度值，成功实现了多重边的分离绘制。",
        "体会：通过本次实验，我掌握了关联矩阵与邻接矩阵的数学转换关系，并学会了使用 Python 库解决图论中的实际可视化问题。"
    ]
    
    for item in summary:
        doc.add_paragraph(item, style='List Bullet')

    # 保存文件
    file_name = '吴虹霖_离散数学II_实验一报告.docx'
    doc.save(file_name)
    print(f"成功生成文件：{file_name}")

if __name__ == "__main__":
    create_experiment_report()