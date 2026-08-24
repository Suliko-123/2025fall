from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_report():
    doc = Document()

    # 设置中文字体支持
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10.5)

    # --- 标题 ---
    title = doc.add_heading('实验报告三', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 个人信息表格  ---
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    data = [
        ['学 号', '24020007135', '姓 名', '吴虹霖'],
        ['专业班级', '24计算机科学与技术', '课程名称', '离散数学II实验'],
        ['学 期', '2025秋', '任课教师', '曹媛'],
        ['完成日期', '2026.01.05', '上机时间', '周二三四节'],
        ['实验名称', '城市交通网络中绕路方案的分析与规划', '', '']
    ]
    
    for i, row_data in enumerate(data):
        row = table.rows[i].cells
        for j, val in enumerate(row_data):
            row[j].text = val

    # 合并最后一行后两格
    table.rows[4].cells[2].merge(table.rows[4].cells[3])

    # --- 一、实验目的 [cite: 34-37, 61] ---
    doc.add_heading('一、实验目的', level=1)
    purposes = [
        "掌握无向连通图生成树和总个数的求解方法 [cite: 35]。",
        "掌握基本回路系统和环路空间的求解方法 [cite: 36]。",
        "了解生成树、环路空间的实际应用（如最小必要路网与绕行规划） [cite: 37, 39]。"
    ]
    for p in purposes:
        doc.add_paragraph(p, style='List Bullet')

    # --- 二、实验要求 [cite: 40-46, 62] ---
    doc.add_heading('二、实验要求', level=1)
    reqs = [
        "用相邻矩阵描述交通网络并进行边标定 [cite: 42]。",
        "构建“最小必要路网”（生成树）并给出其相邻矩阵 [cite: 43]。",
        "计算所有可能的生成树总个数 [cite: 44]。",
        "求解基本回路系统作为“最小绕行单元库” [cite: 45]。",
        "求解环路空间以规划所有绕行方案 [cite: 46]。"
    ]
    for r in reqs:
        doc.add_paragraph(r, style='List Bullet')

    # --- 三、实验内容及步骤 [cite: 63-69] ---
    doc.add_heading('三、实验内容及步骤', level=1)
    
    doc.add_heading('1. 需求分析 [cite: 64]', level=2)
    doc.add_paragraph("输入：无向简单连通图的邻接矩阵 [cite: 65]。\n"
                      "输出：所有生成树及其数量、特定生成树矩阵、基本回路系统及环路空间 [cite: 66, 49-51]。")

    doc.add_heading('2. 概要设计 [cite: 68]', level=2)
    doc.add_paragraph("数据结构：使用并查集（DisjointSet）处理连通性，使用邻接表（adj_list）存储图结构。\n"
                      "逻辑流程：标定边 -> 组合搜索生成树 -> DFS查找树中路径 -> 异或运算生成环路空间。")

    doc.add_heading('3. 详细代码实现（核心部分） [cite: 69]', level=2)
    # 插入一段代码片段
    code_text = (
        "def get_ordered_cycle_string(edge_set, edge_map):\n"
        "    # 核心：将边集合转换为按物理路径顺序排列的字符串\n"
        "    local_adj = {} ... # 构建局部邻接表并沿边遍历\n"
        "    return \"\".join(ordered_edges)"
    )
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code_text)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)

    # --- 四、实验结果分析 [cite: 56, 73] ---
    doc.add_heading('四、实验结果分析', level=1)
    doc.add_paragraph("针对示例数据（5x5矩阵）：")
    doc.add_paragraph("1. 边标定结果：e1:(1,2), e2:(1,3), e3:(1,5), e4:(2,3)...")
    doc.add_paragraph("2. 生成树统计：程序成功识别所有生成树并计算总数 [cite: 49]。")
    doc.add_paragraph("3. 基本回路：成功输出如 {e1e4e2} 等形式的最小绕行单元 [cite: 50]。")

    # --- 五、实验总结 [cite: 57, 74] ---
    doc.add_heading('五、实验总结', level=1)
    summary = [
        "通过并查集优化了生成树的判定效率。",
        "解决了环路输出无序的问题，通过路径排序使结果符合实际交通绕行逻辑。",
        "掌握了如何利用图论算法解决实际的城市交通稳健性分析问题 [cite: 39]。"
    ]
    for s in summary:
        doc.add_paragraph(s, style='List Bullet')

    # 保存文档
    file_name = "19-吴虹霖-第三次实验报告.docx"
    doc.save(file_name)
    print(f"报告已生成: {file_name}")

if __name__ == "__main__":
    create_report()