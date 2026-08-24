import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name, size, is_bold=False):
    """设置中文字体和西文字体的通用函数"""
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = is_bold

def create_report():
    doc = Document()

    # --- 标题 ---
    header = doc.add_heading('', level=1)
    run_title = header.add_run('实验报告二')
    set_font(run_title, '黑体', 18, True)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # [cite_start]--- 个人信息表格 [cite: 2] ---
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    info = [
        ['学 号', '24020007135', '姓 名', '吴虹霖'],
        ['专业班级', '24计算机科学与技术', '课程名称', '离散数学II实验'],
        ['学 期', '2025秋', '任课教师', '曹媛'],
        ['完成日期', '2025.12.26', '上机时间', '周二三四节'],
        ['实验名称', '无人机巡检路径规划（哈密顿回路与TSP问题）', '', '']
    ]
    # 合并最后一行
    table.cell(4, 1).merge(table.cell(4, 3))
    
    for r in range(5):
        for c in range(4):
            if not (r == 4 and c > 1):
                cell = table.cell(r, c)
                cell.text = info[r][c]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_font(run, '宋体', 11)

    doc.add_paragraph() # 空行

    # [cite_start]--- 一、实验目的 [cite: 25-29] ---
    h1 = doc.add_paragraph()
    set_font(h1.add_run('一、实验目的'), '黑体', 14, True)
    purposes = [
        "理解哈密顿图的定义及最短哈密顿回路的基本性质。",
        "实现图和巡检回路的可视化表示。",
        "掌握最短哈密顿回路的精确解和近似解求解算法。",
        "了解NP完全问题的特性及时间复杂度分析。"
    ]
    for p in purposes:
        para = doc.add_paragraph(style='List Bullet')
        set_font(para.add_run(p), '宋体', 11)

    # [cite_start]--- 二、实验要求 [cite: 36-39] ---
    h2 = doc.add_paragraph()
    set_font(h2.add_run('二、实验要求'), '黑体', 14, True)
    requirements = [
        "对无人机巡检问题进行建模，构建无向完全带权图。",
        "对 n≤6 阶图提供精确最优解及其可视化图示。",
        "对 n≥7 阶图提供近似最优解及其可视化图示。"
    ]
    for r in requirements:
        para = doc.add_paragraph(style='List Bullet')
        set_font(para.add_run(r), '宋体', 11)

    # [cite_start]--- 三、实验内容及步骤 [cite: 60-66] ---
    h3 = doc.add_paragraph()
    set_font(h3.add_run('三、实验内容及步骤'), '黑体', 14, True)
    
    # 1. 需求与概要设计
    sub1 = doc.add_paragraph()
    set_font(sub1.add_run('1. 需求与概要设计'), '宋体', 12, True)
    # 此处已修正报错逻辑
    para_logic = doc.add_paragraph("输入：一个 n x n 的邻接矩阵。处理逻辑：利用全排列算法（优化去重）实现精确解，利用最近邻贪心策略实现近似解；通过 Matplotlib 进行路径的可视化展示。")
    for run in para_logic.runs:
        set_font(run, '宋体', 11)

    # 2. 详细代码实现
    sub2 = doc.add_paragraph()
    set_font(sub2.add_run('2. 详细代码实现（核心算法）'), '宋体', 12, True)
    code_text = (
        "def solve_exact_optimized(matrix):\n"
        "    # 核心优化：消除逆时针重复路径\n"
        "    for p in permutations(nodes_to_permute):\n"
        "        if p[0] < p[-1]:\n"
        "            path = [0] + list(p) + [0]\n"
        "            # ...计算总距离并更新最小值...\n"
    )
    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent = Inches(0.3)
    run_code = code_para.add_run(code_text)
    set_font(run_code, 'Consolas', 10)
    run_code.font.color.rgb = RGBColor(0, 102, 204)

    # [cite_start]--- 四、实验结果分析 [cite: 70] ---
    h4 = doc.add_paragraph()
    set_font(h4.add_run('四、实验结果分析'), '黑体', 14, True)
    res_text = (
        "对于测试数据 A1（n=4）：程序计算出精确最优解路径，总长度为 11，证明了全排列搜索的有效性。\n"
        "对于测试数据 A2（n=7）：由于节点增多，程序采用最近邻近似算法。虽然复杂度显著降低（O(n^2)），但仍能提供符合巡检要求的低成本路径。"
    )
    para_res = doc.add_paragraph()
    set_font(para_res.add_run(res_text), '宋体', 11)

    # [cite_start]--- 五、心得总结 [cite: 71] ---
    h5 = doc.add_paragraph()
    set_font(h5.add_run('五、心得总结'), '黑体', 14, True)
    summary = (
        "1. 遇到的问题：在使用全排列算法处理大规模节点时，计算时间呈爆炸式增长。\n"
        "2. 解决办法：通过限制 n 值范围，对规模较大的图引入最近邻近似算法，实现了精度与速度的平衡。\n"
        "3. 心得：本次实验加深了我对 TSP 问题及 NP 完全性特征的理解，掌握了图论算法在实际路径规划中的应用。"
    )
    para_sum = doc.add_paragraph()
    set_font(para_sum.add_run(summary), '宋体', 11)

    # 保存文件
    save_path = '19-吴虹霖-第二次实验报告_修正版.docx'
    doc.save(save_path)
    print(f"报告已成功生成：{save_path}")

if __name__ == "__main__":
    create_report()